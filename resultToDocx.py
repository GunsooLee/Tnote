from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os

def create_meeting_minutes(title, meeting_room, date, writer, attendees, subject, contents, meeting_time, stt, att_suj, file_name):
    """_summary_
    Args:
        title (string): 회의록 문서 제목
        meeting_room (string): 회의실 
        date (string): 회의 날짜
        writer (string): 작성자
        attendees (string): 참석 인원
        contents (string): 회의 내용
        stt (string): stt 결과
        att_suj (strinf): 화자별
        file_name (string): 다운로드 파일 이름
    """

    document = Document()

    # 문서 정보
    document.add_heading('T-note 회의록', 0)
    
    # 회의 제목
    document.add_heading('회의 제목', level=1)
    document.add_paragraph(title)
    
    # 회의정보
    document.add_heading('회의 정보', level=1)
    p = document.add_paragraph()
    p.add_run('일시 : ').bold=True
    p.add_run(date)
    p.add_run('\n')
    p.add_run('회의 시간 : ').bold=True
    p.add_run(meeting_time)
    p.add_run('\n')
    
    # 회의 장소
    p.add_run('장소 : ').bold=True
    p.add_run(meeting_room)
    p.add_run('\n')
    # 회의 참석자
    p.add_run('참석자 : ').bold=True
    p.add_run(", ".join(attendees))
    
    p.add_run('\n')
    # 회의 작성자
    p.add_run('작성자 : ').bold=True
    p.add_run(writer)     

    # 회의 주제
    document.add_heading('회의 주제', level=1)
    document.add_paragraph(subject)

    # 회의 내용
    document.add_heading('회의 내용', level=1)
    document.add_paragraph(contents)

    # 화자별 요약 페이지
    document.add_page_break()    
    document.add_heading('화자별 요약', level=1)
    document.add_paragraph(att_suj)
    
    # STT 결과 페이지
    document.add_page_break()
    document.add_heading('STT 결과', level=1)
    document.add_paragraph(stt)
    
    
    # 파일 저장 
    file_path = 'result_file/' + file_name + '.docx'
    document.save(file_path)
    
    #파일 사이즈
    file_size =  os.path.getsize(file_path)
    file_size = file_size / (1024)
    return file_size, file_path
